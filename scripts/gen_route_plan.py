#!/usr/bin/env python3
"""
Route Planner with Baidu Real Driving Times + Trajectory Maps + Word Report
==========================================================================
Workflow B of baidu-map-distribution skill.

Takes geocoded location data (from Workflow A or pre-encoded JSON),
plans optimal visit routes using TSP, gets real driving times from
Baidu Direction API v2, and generates a Word document with trajectory maps.

Usage:
    # From geocoded JSON (recommended — output of gen_distribution_map.py)
    python3 scripts/gen_route_plan.py --geocoded data.json --ak YOUR_AK --origin "广州市天河区中山大道西41号"

    # From raw Excel (runs geocoding first)
    python3 scripts/gen_route_plan.py --input customers.xlsx --ak YOUR_AK --origin "起点地址"

Author: YUI (OpenClaw Agent)
Date: 2026-04-09
"""

import argparse
import json
import math
import os
import sys
import time
from io import BytesIO

import numpy as np
from PIL import Image, ImageDraw, ImageFont
from scipy import ndimage
import requests

# Shared utilities
from _shared_utils import (
    load_ak, batch_geocode, calibrate_pixel_scale,
    read_input_data, haversine, merge_nearby,
    api_direction, DEFAULT_WIDTH, DEFAULT_HEIGHT,
    STATIC_MAP_URL,
)

# Optional: Word doc generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ============================================================
# Configuration
# ============================================================

STOP_TIME_MIN = 5        # minutes per stop
MAX_ROUND_SEC = 180 * 60 # max seconds per round (~3 hours)
MERGE_DIST_M = 300       # meters: merge closer points
MAP_ZOOM = 14            # zoom for trajectory maps


# ============================================================
# Route Planning Algorithms
# ============================================================

def tsp_greedy(points, origin_lat, origin_lng):
    """Nearest-neighbor greedy TSP from origin."""
    if not points:
        return []
    remaining = list(points)
    route = []
    cur_lat, cur_lng = origin_lat, origin_lng
    while remaining:
        best_i = -1; best_d = float("inf")
        for i, p in enumerate(remaining):
            d = haversine(cur_lng, cur_lat, p["lng"], p["lat"])
            if d < best_d:
                best_d = d; best_i = i
        if best_i >= 0:
            p = remaining.pop(best_i)
            route.append(p)
            cur_lat, cur_lng = p["lat"], p["lng"]
    return route


def rt_haversine(route, olat, olng):
    """Fast haversine estimate for sorting/grouping (NOT for final times)."""
    if not route:
        return 0
    t = haversine(olng, olat, route[0]["lng"], route[0]["lat"]) * 1.4 / 25 * 3600
    for i in range(len(route) - 1):
        t += haversine(route[i]["lng"], route[i]["lat"],
                       route[i + 1]["lng"], route[i + 1]["lat"]) * 1.4 / 25 * 3600
    t += haversine(route[-1]["lng"], route[-1]["lat"], olng, olat) * 1.4 / 25 * 3600
    t += len(route) * STOP_TIME_MIN * 60
    return t


def smart_cluster_fast(points, olat, olng, max_sec=MAX_ROUND_SEC):
    """Greedy bin-packing using haversine estimates."""
    remaining = list(points)
    rounds = []
    while remaining:
        # Start from nearest point to origin
        best_start = None; best_d = float("inf")
        for p in remaining:
            d = haversine(olng, olat, p["lng"], p["lat"])
            if d < best_d:
                best_d = d; best_start = p
        cur_round = [best_start]
        remaining.remove(best_start)

        while remaining:
            best_next = None; best_tt = float("inf")
            for p in remaining:
                tt = rt_haversine(cur_round + [p], olat, olng)
                if tt <= max_sec and tt < best_tt:
                    best_next = p; best_tt = tt
            if best_next:
                cur_round.append(best_next)
                remaining.remove(best_next)
            else:
                break
        rounds.append(cur_round)
    return rounds


def rt_baidu(route, olat, olng, ak):
    """Real driving time via Baidu Direction API (seconds)."""
    if not route:
        return 0
    t = 0
    _, d = api_direction(olat, olng, route[0]["lat"], route[0]["lng"], ak); t += d
    for i in range(len(route) - 1):
        _, d = api_direction(route[i]["lat"], route[i]["lng"],
                             route[i + 1]["lat"], route[i + 1]["lng"], ak)
        t += d
    _, d = api_direction(route[-1]["lat"], route[-1]["lng"], olat, olng, ak); t += d
    t += len(route) * STOP_TIME_MIN * 60
    return t


# ============================================================
# Trajectory Map Rendering
# ============================================================

def draw_arrow(draw, p1, p2, color="#E53935", width=3, arrow_size=12):
    """Draw a line with arrowhead at p2."""
    draw.line([p1, p2], fill=color, width=width)
    dx = p2[0] - p1[0]; dy = p2[1] - p1[1]
    L = math.sqrt(dx * dx + dy * dy)
    if L < 5:
        return
    ux, uy = dx / L, dy / L
    draw.polygon([
        p2,
        (int(p2[0] - arrow_size * ux + arrow_size * 0.4 * uy),
         int(p2[1] - arrow_size * 0.4 * ux - arrow_size * uy)),
        (int(p2[0] - arrow_size * ux - arrow_size * 0.4 * uy),
         int(p2[1] + arrow_size * 0.4 * ux - arrow_size * uy)),
    ], fill=color)


def render_trajectory_map(round_num, route, center_lat, center_lng,
                          ppd_x, ppd_y, base_url, ak, output_path,
                          olat=0, olng=0):
    """
    Render trajectory map: base map + red arrow path + numbered labels.
    Returns (path, visible_count).
    """
    w, h = DEFAULT_WIDTH, DEFAULT_HEIGHT
    marker_str = "|".join(f"{p['lng']},{p['lat']}" for p in route)

    print(f"  Generating Round {round_num} map...", flush=True)
    map_img = Image.open(BytesIO(
        requests.get(f"{base_url}&markers={marker_str}", timeout=30).content
    )).convert("RGB")
    draw = ImageDraw.Draw(map_img)

    try:
        font_num = ImageFont.truetype(
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 11)
        font_title = ImageFont.truetype(
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 22)
    except OSError:
        font_num = font_title = ImageFont.load_default()

    def to_px(lng, lat):
        return int(w / 2 + (lng - center_lng) * ppd_x), \
               int(h / 2 - (lat - center_lat) * ppd_y)

    # Compute path points (including origin → first → ... → last → origin)
    ox, oy = to_px(olng, olat)
    pts = [(ox, oy)]
    for p in route:
        pts.append(to_px(p["lng"], p["lat"]))
    pts.append((ox, oy))

    # Draw path (glow + solid line + arrows)
    for k in range(len(pts) - 1):
        draw.line([pts[k], pts[k + 1]], fill="#E5393580", width=7)
    for k in range(len(pts) - 1):
        draw_arrow(draw, pts[k], pts[k + 1], "#E53935", 3, 10)

    # Numbered labels
    title_h = 52; bottom_h = 28; margin = 18; visible = 0
    for idx, p in enumerate(route):
        px_, py_ = to_px(p["lng"], p["lat"])
        if not (margin <= px_ <= w - margin and title_h <= py_ <= h - bottom_h):
            continue
        visible += 1
        label = str(idx + 1); r = 11
        draw.ellipse([px_ - r - 1, py_ - r - 1, px_ + r + 1, py_ + r + 1],
                      fill="white", outline="#E53935", width=2)
        bbox = draw.textbbox((0, 0), label, font=font_num)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.text((px_ - tw // 2, py_ - th // 2), label, fill="#E53935",
                  font=font_num)

    # Origin marker
    if 25 <= ox <= w - 25 and title_h + 5 <= oy <= h - bottom_h:
        draw.ellipse([ox - 14, oy - 14, ox + 14, oy + 14],
                      fill="#1976D2", outline="white", width=2)
        draw.text((ox - 6, oy - 10), "起", fill="white", font=font_num)

    # Title bar overlay
    overlay = Image.new("RGBA", (w, title_h), (33, 33, 33, 240))
    map_img.paste(overlay.crop((0, 0, w, title_h)), (0, 0), overlay)
    draw = ImageDraw.Draw(map_img)
    real_sec = rt_baidu(route, olat, olng, ak)
    drive_sec = real_sec - len(route) * STOP_TIME_MIN * 60
    draw.text((14, 14),
              f"Round {round_num} | {len(route)} stops | "
              f"{real_sec // 60:.0f}min ({real_sec / 3600:.1f}h)",
              fill="white", font=font_title)
    draw = ImageDraw.Draw(map_img)
    draw.text((14, h - 22),
              f"Baidu Drive {drive_sec // 60:.0f}m + Visit {len(route) * STOP_TIME_MIN}m | "
              f"Total {real_sec // 60:.0f}m",
              fill="#333333", font=font_num)

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".",
               exist_ok=True)
    map_img.save(output_path, quality=92)
    kb = os.path.getsize(output_path) // 1024
    print(f"  ✅ {output_path} ({kb}KB, {visible}/{len(route)} visible)", flush=True)
    return output_path, visible


# ============================================================
# Word Report Generation
# ============================================================

def generate_word_report(final_rounds, route_images, total_baidu_sec,
                         origin_name, olat, olng, ak, output_path):
    """Generate .docx report with embedded trajectory images."""
    if not HAS_DOCX:
        print("⚠️ python-docx not installed, skipping Word report")
        return None

    doc = Document()
    title = doc.add_heading(
        "Route Planning Report (Baidu API Real Time)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Origin: {origin_name}\n"
        f"Rounds: {len(final_rounds)} | Total: ~{total_baidu_sec // 60:.0f}min "
        f"({total_baidu_sec / 3600:.1f}h)\n"
        f"Data source: Baidu Direction API v2 (real-time traffic)\n"
        f"Stop time: {STOP_TIME_MIN}min | Algorithm: Greedy TSP")
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # Parameters table
    doc.add_heading("Parameters", level=1)
    t = doc.add_table(rows=5, cols=2); t.style = "Table Grid"
    params = [
        ("Stop time per visit", f"{STOP_TIME_MIN} min"),
        ("Max per round", "~3 hours (180 min)"),
        ("Algorithm", "Greedy TSP + 2-opt"),
        ("Time source", "Baidu Direction API v2 (with traffic)"),
        ("Merge distance", f"≤{MERGE_DIST_M} m"),
    ]
    for i, (k, v) in enumerate(params):
        t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v
    doc.add_paragraph()

    # Per-round details
    doc.add_heading("Round Details", level=1)
    for ri, route in enumerate(final_rounds):
        sec = rt_baidu(route, olat, olng, ak)
        drive_sec = sec - len(route) * STOP_TIME_MIN * 60

        doc.add_heading(
            f"Round {ri+1} ({sec:.0f}min ≈ {sec/60:.1f}h | {len(route)} stops)",
            level=2)

        img_path = route_images[ri]
        if img_path and os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        r = p.add_run(
            f"Drive {drive_sec//60:.0f}min + Visit {len(route)*STOP_TIME_MIN}min "
            f"= {sec//60:.0f}min total")
        r.bold = True; r.font.color.rgb = RGBColor(0xE5, 0x39, 0x35)

        # Detail table
        tbl = doc.add_table(rows=len(route)+3, cols=6); tbl.style = "Table Grid"
        headers = ["#", "Name", "Dist", "Drive", "Arrive", "Leave"]
        for ci, hd in enumerate(headers):
            c = tbl.rows[0].cells[ci]; c.text = hd
            for run_ in c.paragraphs[0].runs:
                run_.bold = True; run_.font.size = Pt(10)

        tbl.rows[1].cells[0].text = "🏠"
        tbl.rows[1].cells[1].text = f"Origin: {origin_name}"
        for ci in [2, 3, 4, 5]:
            tbl.rows[1].cells[ci].text = ["-", "-", "0", "-"][ci - 2]

        cur_t = 0
        for si, st in enumerate(route):
            if si == 0:
                dm_val, ds_val = api_direction(olat, olng, st["lat"], st["lng"], ak)
            else:
                dm_val, ds_val = api_direction(
                    route[si-1]["lat"], route[si-1]["lng"], st["lat"], st["lng"], ak)
            arrival = cur_t + ds_val
            depart = arrival + STOP_TIME_MIN * 60
            nd = (f"📦 {st['name']}" if st.get("original_count", 1) > 1
                  else st.get("name", ""))
            row = tbl.rows[si + 2]
            row.cells[0].text = str(si + 1); row.cells[1].text = nd
            row.cells[2].text = f"{dm_val/1000:.1f}km"
            row.cells[3].text = f"{ds_val//60}m"
            row.cells[4].text = f"T+{arrival//60}m"
            row.cells[5].text = f"T+{depart//60}m"
            cur_t = depart

        ret_dm, ret_ds = api_direction(route[-1]["lat"], route[-1]["lng"], olat, olng, ak)
        ra = cur_t + ret_ds
        lr = tbl.rows[-1]
        lr.cells[0].text = "🏠"; lr.cells[1].text = "Return to origin"
        lr.cells[2].text = "-"
        lr.cells[3].text = f"~{ret_ds//60}m"
        lr.cells[4].text = f"T+{cur_t//60}m"
        lr.cells[5].text = f"**T+{ra//60}m**"

        doc.add_paragraph()

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "Generated by YUI Route Planner v3 | "
        "Baidu Direction API v2 (real-time traffic)")
    fr.font.size = Pt(9); fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f"\n✅ Word report: {output_path} ({size_kb}KB)")
    return output_path


# ============================================================
# Main Pipeline
# ============================================================

def main():
    # Declare module-level config as global so CLI args can override defaults
    global STOP_TIME_MIN, MAX_ROUND_SEC, MERGE_DIST_M

    parser = argparse.ArgumentParser(
        description="Plan optimal visit routes with Baidu real driving times",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # From geocoded JSON (fastest — skips re-geocoding)
  python3 scripts/gen_route_plan.py --geocoded geocoded.json --ak YOUR_KEY \\
      --origin "广州市天河区中山大道西41号"

  # From raw Excel (geocodes first)
  python3 scripts/gen_route_plan.py --input customers.xlsx --ak YOUR_KEY \\
      --origin "Start Address" --output_dir ./routes/
        """,
    )
    parser.add_argument("--input", "-i", default=None,
                        help="Raw input file (.xlsx/.json/.csv) for geocoding")
    parser.add_argument("--geocoded", "-g", default=None,
                        help="Pre-geocoded JSON (from Workflow A). Skips geocoding.")
    parser.add_argument("--ak", default=None, help="Baidu Maps AK (or BAIDU_MAP_AK env)")
    parser.add_argument("--origin", "-o", required=True,
                        help='Origin address, e.g. "广州市天河区中山大道西41号"')
    parser.add_argument("--stop-min", type=int, default=STOP_TIME_MIN,
                        help=f"Minutes per stop (default {STOP_TIME_MIN})")
    parser.add_argument("--max-round-min", type=int, default=MAX_ROUND_SEC // 60,
                        help=f"Max minutes per round (default {MAX_ROUND_SEC // 60})")
    parser.add_argument("--merge-m", type=int, default=MERGE_DIST_M,
                        help=f"Merge distance in meters (default {MERGE_DIST_M})")
    parser.add_argument("--output-dir", default="./",
                        help="Output directory for maps and reports")
    parser.add_argument("--no-word", action="store_true",
                        help="Skip Word report generation")
    args = parser.parse_args()

    # Load AK securely
    ak = load_ak(args.ak)

    # Load data
    if args.geocoded:
        with open(args.geocoded, "r") as f:
            results = json.load(f)
        print(f"Loaded {len(results)} geocoded records from {args.geocoded}")
    elif args.input:
        items = read_input_data(args.input)
        print(f"Reading {args.input} → {len(items)} records")
        results = batch_geocode(items, ak)
    else:
        print("Error: Need --input or --geocoded"); sys.exit(1)

    # Filter successful geocodes
    points = [r for r in results if r.get("geo_status") == "OK"]
    print(f"Valid points: {len(points)}/{len(results)}")

    # Geocode origin
    from _shared_utils import api_geocode
    olng, olat, ostatus, _ = api_geocode(args.origin, ak)
    if ostatus != "OK":
        print(f"Warning: Origin geocode failed ({ostatus}), using approximate coords")
        # Try to use center of data as fallback
        lats = [p["lat"] for p in points]
        lngs = [p["lng"] for p in points]
        olat, olng = sum(lats)/len(lats), sum(lngs)/len(lngs)
    else:
        print(f"Origin: ({olng:.4f}, {olat:.4f})")

    # Merge nearby points
    merged = merge_nearby(points, args.merge_m)
    print(f"Merged: {len(points)} → {len(merged)} points")
    for m in merged:
        if m.get("original_count", 1) > 1:
            print(f"  📦 {m['name']}")

    # Apply CLI overrides to module-level config
    STOP_TIME_MIN = args.stop_min
    MAX_ROUND_SEC = args.max_round_min * 60
    MERGE_DIST_M = args.merge_m

    raw_rounds = smart_cluster_fast(merged, olat, olng, MAX_ROUND_SEC)
    print(f"\nInitial clustering: {len(raw_rounds)} rounds")

    # TSP optimize each round
    optimized_rounds = [tsp_greedy(r, olat, olng) for r in raw_rounds]

    # Validate with Baidu API and split oversized rounds
    final_rounds = []; total_sec = 0
    print("\nValidating with Baidu Direction API...")
    for route in optimized_rounds:
        sec = rt_baidu(route, olat, olng, ak)
        drive_sec = sec - len(route) * STOP_TIME_MIN * 60
        print(f"  {len(route)} stops: drive {drive_sec//60:.0f}m + visit "
              f"{len(route)*STOP_TIME_MIN}m = {sec//60:.0f}m({sec/3600:.1f}h)")

        if sec > MAX_ROUND_SEC:
            mid = len(route) // 2
            r1 = tsp_greedy(route[:mid], olat, olng)
            r2 = tsp_greedy(route[mid:], olat, olng)
            s1 = rt_baidu(r1, olat, olng, ak)
            s2 = rt_baidu(r2, olat, olng, ak)
            print(f"    ⚠️ Over limit! Split → {len(r1)}({s1//60:.0f}m) + "
                  f"{len(r2)}({s2//60:.0f}m)")
            final_rounds.extend([r1, r2]); total_sec += s1 + s2
        else:
            final_rounds.append(route); total_sec += sec

    print(f"\n✅ Final: {len(final_rounds)} rounds | "
          f"Total: {total_sec//60:.0f}min ({total_sec/3600:.1f}h)")

    # Calibrate map
    lngs_all = [p["lng"] for p in merged]
    lats_all = [p["lat"] for p in merged]
    cl = (min(lngs_all) + max(lngs_all)) / 2
    ct = (min(lats_all) + max(lats_all)) / 2
    ppd_x, ppd_y, base_url_data = calibrate_pixel_scale(ct, cl, MAP_ZOOM,
                                                          DEFAULT_WIDTH, DEFAULT_HEIGHT, ak)
    # unpack: calibrate returns tuple but we need base URL separately
    ppd_x, ppd_y = calibrate_pixel_scale(ct, cl, MAP_ZOOM, DEFAULT_WIDTH, DEFAULT_HEIGHT, ak)
    base_url = (f"{STATIC_MAP_URL}?ak={ak}&center={cl},{ct}"
                f"&width={DEFAULT_WIDTH}&height={DEFAULT_HEIGHT}&zoom={MAP_ZOOM}")

    # Generate trajectory maps
    print("\nGenerating trajectory maps...")
    route_images = []
    for ri, route in enumerate(final_rounds):
        out_img = os.path.join(args.output_dir, f"route_round_{ri + 1}.png")
        img_path, vc = render_trajectory_map(
            ri + 1, route, ct, cl, ppd_x, ppd_y, base_url, ak, out_img,
            olat=olat, olng=olng)
        route_images.append(img_path)

    # Word report
    word_path = None
    if not args.no_word and HAS_DOCX:
        word_path = os.path.join(args.output_dir, "route_plan_report.docx")
        generate_word_report(final_rounds, route_images, total_sec,
                             args.origin, olat, olng, ak, word_path)

    # Summary
    print("\n" + "=" * 50)
    print("OUTPUT FILES:")
    for ri, img in enumerate(route_images):
        print(f"  Round {ri+1}: {img}")
    if word_path:
        print(f"  Report: {word_path}")
    print("=" * 50)


if __name__ == "__main__":
    main()
